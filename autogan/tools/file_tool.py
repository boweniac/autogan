import os
import pandas as pd
import PyPDF2
from typing import Optional
from docx import Document


class File:
    def __init__(self, work_dir: Optional[str] = None):
        """文件操作类

        Read content: word, excel, pdf

        Append content: word, text

        :param work_dir: The relative path of the document, default is extensions
        """
        if work_dir is None:
            work_dir = "extensions"
        self._work_dir = os.getcwd() + "/" + work_dir

    def read(self, file_name):
        """Read content: word, excel, pdf
        """
        file_type = self.get_file_type(file_name)
        if file_type == '.docx':
            result = self.read_word(file_name)
        elif file_type == '.xlsx':
            result = self.read_excel(file_name)
        elif file_type == '.pdf':
            result = self.read_pdf(file_name)
        else:
            result = None

        return result

    @staticmethod
    def get_file_type(file_name) -> str:
        """Get file type
        """
        _, ext = os.path.splitext(file_name)
        return ext

    def read_word(self, file_name):
        try:
            file_path = self._work_dir + "/" + file_name
            text = ""
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except:
            return None

    def read_excel(self, file_name):
        try:
            file_path = self._work_dir + "/" + file_name
            df = pd.read_excel(file_path)
            return df.to_string()
        except:
            return None

    def read_pdf(self, file_name):
        try:
            file_path = self._work_dir + "/" + file_name
            pdf_file_obj = open(file_path, 'rb')
            pdf_reader = PyPDF2.PdfReader(pdf_file_obj)

            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()

            return text
        except:
            return None

    def append_to_text(self, file_name, text) -> bool:
        """Append text to the end of the text file

        :param file_name: text file name
        :param text: Text to be appended

        :return: Successful or not
        """
        try:
            file_path = self._work_dir + "/" + file_name + '.txt'
            with open(file_path, 'a') as file:
                file.write(text)
            return True
        except FileNotFoundError:
            return False

    def append_to_word(self, file_name, text) -> bool:
        """Append text to the end of the word file

        :param file_name: word file name
        :param text: Text to be appended

        :return: Successful or not
        """
        try:
            file_path = self._work_dir + "/" + file_name + '.docx'
            if os.path.exists(file_path):
                doc = Document(file_path)
            else:
                doc = Document()
            doc.add_paragraph(text)
            doc.save(file_path)
            return True
        except Exception as e:
            print(e)
            return False
